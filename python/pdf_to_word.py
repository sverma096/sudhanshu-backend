import sys
import pdfplumber
from docx import Document
from utils import temp_output, print_result

file_path = sys.argv[1]
doc = Document()

with pdfplumber.open(file_path) as pdf:
    for page in pdf.pages:
        text = page.extract_text() or ""
        if text.strip():
            for para in text.split("\n"):
                doc.add_paragraph(para)
        else:
            doc.add_paragraph("")

out = temp_output(".docx")
doc.save(out)
print_result(out, "output.docx")
